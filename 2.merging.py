#Python sript to merge .docx and .pdf format files in a given directory
#do not use special characters or blank spaces to name your files, otherwise the output is error
#however if your source consist of plenty of files, run renamescript.py to normalize their names
#prior to successfully run this code, run a renamescript.py that is going to edit file names in a given directory

#put all files you want to combine in one directory
#process separated in 2 steps - separately merge .docx files and separtely merge .pdf files
#the output are 2 files in a subdirectory of your folder with files

import docx
import time
import os
import Tkinter, tkFileDialog
from PyPDF2 import PdfFileMerger,PdfFileReader

start = time.time()

root = Tkinter.Tk()
root.withdraw()
dirname = tkFileDialog.askdirectory(parent=root,initialdir="/",title='Please select a directory with your files')
os.chdir(dirname)
print os.getcwd()

print '\n'

files = [f for f in os.listdir(dirname) if f.endswith('.docx')]
print files

#create subdirectory to store combined(merged) files created by the following process
outputdir = '.\combinedfiles'
if not os.path.exists(outputdir):
    os.makedirs(outputdir)

output = '.\combinedfiles\combinedocs.docx' #define a name of the file that contains a merging of all .docx files

input = files

def text_from_docx(path):

    doc = docx.Document(path)
    for paragraph in doc.paragraphs:
        yield paragraph.text

combined_doc = docx.Document()
for file_path in input:
    for text in text_from_docx(file_path):
        combined_doc.add_paragraph(text)
        combined_doc.save(output)

#Script to merge pdf

pdfs = [f for f in os.listdir(dirname) if f.endswith('.pdf')]

merger = PdfFileMerger()

for filename in pdfs:
    merger.append(PdfFileReader(os.path.join(dirname,filename),'rb'))

merger.write('.\combinedfiles\combinedpdfs.pdf') #must be not indented

print "Merging all documents in directory took: ", time.time() - start, "sec"